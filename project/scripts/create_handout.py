#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""创建AI提示词培训讲义Word文档"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_code_block(doc, lines):
    """添加代码块"""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        # 设置背景色
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F5F5F5')
        p._p.get_or_add_pPr().append(shading)

def create_handout():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题
    title = doc.add_heading('AI助手进课堂：小学教师提示词入门', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('培训讲义', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 一、培训概述
    doc.add_heading('一、培训概述', 1)
    doc.add_heading('培训目标', 2)
    doc.add_paragraph('通过本次培训，您将：')
    doc.add_paragraph('理解AI是什么，消除陌生感', style='List Bullet')
    doc.add_paragraph('掌握5个核心提示词技巧', style='List Bullet')
    doc.add_paragraph('获得3个可直接使用的提示词模板', style='List Bullet')
    doc.add_paragraph('能够独立使用AI辅助备课', style='List Bullet')
    
    doc.add_heading('培训流程', 2)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    
    # 表头
    headers = ['环节', '时长', '内容']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'E8F4F8')
        cell.paragraphs[0].runs[0].bold = True
    
    # 表格内容
    data = [
        ['开场：神奇时刻', '15分钟', 'AI是什么 + 现场演示'],
        ['案例1：生成教案', '25分钟', '角色设定+任务描述'],
        ['案例2：设计活动', '25分钟', '提供示例+迭代优化'],
        ['案例3：制作课件', '25分钟', '结构化输出要求'],
        ['动手练习', '20分钟', '老师实操+答疑'],
        ['总结收尾', '10分钟', '模板发放+学习建议'],
    ]
    for i, row_data in enumerate(data):
        for j, text in enumerate(row_data):
            table.rows[i+1].cells[j].text = text
    
    # 二、AI基础知识
    doc.add_heading('二、AI基础知识', 1)
    doc.add_heading('AI是什么？', 2)
    doc.add_paragraph('AI（人工智能）就像一个「超级助手」：')
    doc.add_paragraph('它需要你给它清晰的指令', style='List Bullet')
    doc.add_paragraph('提示词就是给助手的「工作说明」', style='List Bullet')
    doc.add_paragraph('它不会取代老师，但会用提示词的老师效率会高10倍', style='List Bullet')
    
    doc.add_heading('推荐AI工具', 2)
    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Table Grid'
    
    headers2 = ['工具名称', '特点', '适用场景']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'E8F4F8')
        cell.paragraphs[0].runs[0].bold = True
    
    data2 = [
        ['Kimi', '免费、支持长文本', '教案、课件内容生成'],
        ['文心一言', '百度出品、中文优化', '各类教学场景'],
        ['豆包', '字节出品、界面友好', '新手入门'],
        ['通义千问', '阿里出品、功能全面', '各类教学场景'],
    ]
    for i, row_data in enumerate(data2):
        for j, text in enumerate(row_data):
            table2.rows[i+1].cells[j].text = text
    
    doc.add_heading('使用方式', 2)
    doc.add_paragraph('打开AI工具网站或APP', style='List Number')
    doc.add_paragraph('在对话框中输入提示词', style='List Number')
    doc.add_paragraph('等待AI生成内容', style='List Number')
    doc.add_paragraph('根据需要修改或追问', style='List Number')
    
    # 三、5个核心提示词技巧
    doc.add_heading('三、5个核心提示词技巧', 1)
    
    # 技巧1
    doc.add_heading('技巧1：角色设定', 2)
    p = doc.add_paragraph()
    p.add_run('原理：').bold = True
    p.add_run('告诉AI它是谁，它会以这个身份来思考和回答')
    
    p = doc.add_paragraph()
    p.add_run('示例：').bold = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('❌ ').font.color.rgb = RGBColor(255, 0, 0)
    p.add_run('「帮我写一个教案」')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('✅ ').font.color.rgb = RGBColor(0, 170, 0)
    p.add_run('「你是一位有20年经验的小学语文教师，擅长情境教学...」')
    
    p = doc.add_paragraph()
    p.add_run('效果：').bold = True
    p.add_run('AI会以资深教师的视角来设计，更符合教学实际')
    
    # 技巧2
    doc.add_heading('技巧2：任务描述', 2)
    p = doc.add_paragraph()
    p.add_run('原理：').bold = True
    p.add_run('明确告诉AI你要什么，越具体越好')
    
    p = doc.add_paragraph()
    p.add_run('关键要素：').bold = True
    doc.add_paragraph('学科、年级', style='List Bullet')
    doc.add_paragraph('课题名称', style='List Bullet')
    doc.add_paragraph('课时数量', style='List Bullet')
    doc.add_paragraph('具体要求', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('示例：').bold = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('❌ ').font.color.rgb = RGBColor(255, 0, 0)
    p.add_run('「写一个教案」')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('✅ ').font.color.rgb = RGBColor(0, 170, 0)
    p.add_run('「设计一节三年级语文《荷花》第一课时的教案」')
    
    # 技巧3
    doc.add_heading('技巧3：提供示例', 2)
    p = doc.add_paragraph()
    p.add_run('原理：').bold = True
    p.add_run('给AI看一个你喜欢的范例，它会模仿这个风格')
    
    p = doc.add_paragraph()
    p.add_run('示例：').bold = True
    add_code_block(doc, [
        '参考风格：我希望活动像「知识抢答赛」那样，',
        '学生分组竞争，气氛活跃'
    ])
    
    # 技巧4
    doc.add_heading('技巧4：迭代优化', 2)
    p = doc.add_paragraph()
    p.add_run('原理：').bold = True
    p.add_run('第一次生成不满意时，提出具体修改要求')
    
    p = doc.add_paragraph()
    p.add_run('常用追问方式：').bold = True
    doc.add_paragraph('「请把教学目标写得更具体一些」', style='List Bullet')
    doc.add_paragraph('「活动时间太长了，请缩短到10分钟」', style='List Bullet')
    doc.add_paragraph('「请增加一个互动环节」', style='List Bullet')
    
    # 技巧5
    doc.add_heading('技巧5：结构化输出要求', 2)
    p = doc.add_paragraph()
    p.add_run('原理：').bold = True
    p.add_run('明确告诉AI你要什么格式')
    
    p = doc.add_paragraph()
    p.add_run('示例：').bold = True
    add_code_block(doc, [
        '请按以下格式输出：',
        '1. 活动名称',
        '2. 活动规则',
        '3. 所需材料',
        '4. 教师引导语'
    ])
    
    # 四、提示词模板
    doc.add_heading('四、提示词模板', 1)
    
    # 模板1
    doc.add_heading('模板1：教案生成', 2)
    add_code_block(doc, [
        '你是一位有20年经验的小学【学科】教师，擅长【教学特色】。',
        '',
        '请帮我设计一节【年级】【学科】课的教案框架：',
        '- 课题：【课题名称】',
        '- 课时：【X】课时',
        '- 教学目标需要包含：知识与技能、过程与方法、情感态度价值观',
        '- 请提供：教学目标、教学重难点、教学过程（含导入、新授、练习、总结）、板书设计'
    ])
    
    p = doc.add_paragraph()
    p.add_run('使用示例：').bold = True
    add_code_block(doc, [
        '你是一位有20年经验的小学语文教师，擅长情境教学。',
        '',
        '请帮我设计一节三年级语文课的教案框架：',
        '- 课题：《荷花》',
        '- 课时：2课时',
        '- 教学目标需要包含：知识与技能、过程与方法、情感态度价值观',
        '- 请提供：教学目标、教学重难点、教学过程（含导入、新授、练习、总结）、板书设计'
    ])
    
    # 模板2
    doc.add_heading('模板2：活动设计', 2)
    add_code_block(doc, [
        '你是一位擅长游戏化教学的小学【学科】教师。',
        '',
        '请帮我设计一个【年级】【课题】的课堂活动：',
        '- 活动时长：【X】分钟',
        '- 活动目标：【具体目标】',
        '- 学生人数：【X】人',
        '',
        '参考风格：我希望活动像【示例描述，如「抢答游戏」「角色扮演」「小组竞赛」】',
        '',
        '请提供：',
        '1. 活动名称',
        '2. 活动规则（学生能听懂的表述）',
        '3. 所需材料',
        '4. 教师引导语'
    ])
    
    p = doc.add_paragraph()
    p.add_run('使用示例：').bold = True
    add_code_block(doc, [
        '你是一位擅长游戏化教学的小学数学教师。',
        '',
        '请帮我设计一个四年级《分数的初步认识》的课堂活动：',
        '- 活动时长：15分钟',
        '- 活动目标：让学生理解分数的含义',
        '- 学生人数：40人',
        '',
        '参考风格：我希望活动像「小组竞赛」，学生分组动手操作',
        '',
        '请提供：',
        '1. 活动名称',
        '2. 活动规则（学生能听懂的表述）',
        '3. 所需材料',
        '4. 教师引导语'
    ])
    
    # 模板3
    doc.add_heading('模板3：课件内容', 2)
    add_code_block(doc, [
        '你是一位擅长制作教学课件的小学【学科】教师。',
        '',
        '请帮我设计一个【年级】【课题】的PPT课件内容大纲：',
        '- 课件页数：约【X】页',
        '- 教学重点：【重点内容】',
        '',
        '请按以下格式输出每一页：',
        '',
        '【第X页】标题：______',
        '- 核心内容（3-5个要点）',
        '- 配图建议',
        '- 讲解提示（教师可以说的话）',
        '',
        '要求：',
        '1. 每页内容不超过5个要点',
        '2. 语言适合小学生理解',
        '3. 包含1-2个互动问题'
    ])
    
    p = doc.add_paragraph()
    p.add_run('使用示例：').bold = True
    add_code_block(doc, [
        '你是一位擅长制作教学课件的小学科学教师。',
        '',
        '请帮我设计一个五年级《地球的运动》的PPT课件内容大纲：',
        '- 课件页数：约10页',
        '- 教学重点：地球自转和公转的特点',
        '',
        '请按以下格式输出每一页：',
        '',
        '【第X页】标题：______',
        '- 核心内容（3-5个要点）',
        '- 配图建议',
        '- 讲解提示（教师可以说的话）',
        '',
        '要求：',
        '1. 每页内容不超过5个要点',
        '2. 语言适合小学生理解',
        '3. 包含1-2个互动问题'
    ])
    
    # 五、常见问题解答
    doc.add_heading('五、常见问题解答', 1)
    
    doc.add_heading('Q1：AI生成的内容太泛泛，怎么办？', 3)
    p = doc.add_paragraph()
    p.add_run('原因：').bold = True
    p.add_run('提示词不够具体')
    p = doc.add_paragraph()
    p.add_run('解决：').bold = True
    p.add_run('补充更多背景信息')
    doc.add_paragraph('添加学生情况描述', style='List Bullet')
    doc.add_paragraph('说明教学重点难点', style='List Bullet')
    doc.add_paragraph('提供教材版本信息', style='List Bullet')
    
    doc.add_heading('Q2：AI生成的内容太长，怎么办？', 3)
    p = doc.add_paragraph()
    p.add_run('解决：').bold = True
    p.add_run('在提示词中添加字数限制')
    add_code_block(doc, [
        '请控制在500字以内',
        '请生成3个要点即可'
    ])
    
    doc.add_heading('Q3：AI生成的内容不符合教学实际，怎么办？', 3)
    p = doc.add_paragraph()
    p.add_run('解决：').bold = True
    p.add_run('使用迭代优化')
    add_code_block(doc, [
        '请把导入环节改成游戏形式',
        '请增加学生互动的环节',
        '请把语言改得更适合三年级学生'
    ])
    
    doc.add_heading('Q4：不知道用什么AI工具？', 3)
    p = doc.add_paragraph()
    p.add_run('推荐：').bold = True
    doc.add_paragraph('新手入门：豆包（界面友好）', style='List Bullet')
    doc.add_paragraph('教案课件：Kimi（支持长文本）', style='List Bullet')
    doc.add_paragraph('日常使用：文心一言、通义千问', style='List Bullet')
    
    # 六、后续学习建议
    doc.add_heading('六、后续学习建议', 1)
    
    doc.add_heading('多用', 2)
    doc.add_paragraph('每天尝试用AI完成一个小任务', style='List Bullet')
    doc.add_paragraph('从简单任务开始，逐步增加难度', style='List Bullet')
    
    doc.add_heading('多改', 2)
    doc.add_paragraph('对生成结果不满意时，尝试修改提示词', style='List Bullet')
    doc.add_paragraph('记录有效的提示词，形成自己的模板库', style='List Bullet')
    
    doc.add_heading('多交流', 2)
    doc.add_paragraph('和同事分享好用的提示词', style='List Bullet')
    doc.add_paragraph('组建学习小组，互相学习', style='List Bullet')
    
    # 七、结束语
    doc.add_heading('七、结束语', 1)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('AI不会取代老师，但会用AI的老师会更轻松')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(46, 117, 182)
    
    doc.add_paragraph('希望各位老师通过本次培训，能够：')
    doc.add_paragraph('消除对AI的恐惧感', style='List Number')
    doc.add_paragraph('掌握基本的提示词技巧', style='List Number')
    doc.add_paragraph('在日常备课中尝试使用AI', style='List Number')
    
    p = doc.add_paragraph()
    p.add_run('祝各位老师工作顺利！').bold = True
    
    # 保存文档
    output_path = 'docs/training/ai-prompt-training/training-handout.docx'
    doc.save(output_path)
    print(f'培训讲义Word文档已创建：{output_path}')

if __name__ == '__main__':
    create_handout()
