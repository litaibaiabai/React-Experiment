#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""创建AI提示词模板卡片Word文档"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_code_block(doc, lines):
    """添加代码块"""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(11)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F5F5F5')
        p._p.get_or_add_pPr().append(shading)

def create_templates():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题
    title = doc.add_heading('AI提示词模板卡片', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('可打印卡片，方便日常使用').italic = True
    
    doc.add_paragraph()
    
    # 模板1
    doc.add_heading('模板1：教案生成', 1)
    add_code_block(doc, [
        '你是一位有20年经验的小学【学科】教师，擅长【教学特色】。',
        '',
        '请帮我设计一节【年级】【学科】课的教案框架：',
        '- 课题：【课题名称】',
        '- 课时：【X】课时',
        '- 教学目标需要包含：知识与技能、过程与方法、情感态度价值观',
        '- 请提供：教学目标、教学重难点、教学过程（含导入、新授、练习、总结）、板书设计'
    ])
    
    p = doc.add_paragraph()
    p.add_run('填空说明：').bold = True
    doc.add_paragraph('【学科】：语文/数学/英语/科学...', style='List Bullet')
    doc.add_paragraph('【教学特色】：情境教学/游戏化教学/探究式教学...', style='List Bullet')
    doc.add_paragraph('【年级】：一年级/二年级/三年级...', style='List Bullet')
    doc.add_paragraph('【课题名称】：具体课文或知识点', style='List Bullet')
    doc.add_paragraph('【X】：课时数量', style='List Bullet')
    
    # 分页
    doc.add_page_break()
    
    # 模板2
    doc.add_heading('模板2：活动设计', 1)
    add_code_block(doc, [
        '你是一位擅长游戏化教学的小学【学科】教师。',
        '',
        '请帮我设计一个【年级】【课题】的课堂活动：',
        '- 活动时长：【X】分钟',
        '- 活动目标：【具体目标】',
        '- 学生人数：【X】人',
        '',
        '参考风格：我希望活动像【示例描述】',
        '',
        '请提供：',
        '1. 活动名称',
        '2. 活动规则（学生能听懂的表述）',
        '3. 所需材料',
        '4. 教师引导语'
    ])
    
    p = doc.add_paragraph()
    p.add_run('填空说明：').bold = True
    doc.add_paragraph('【示例描述】：抢答游戏/角色扮演/小组竞赛/动手操作...', style='List Bullet')
    
    # 分页
    doc.add_page_break()
    
    # 模板3
    doc.add_heading('模板3：课件内容', 1)
    add_code_block(doc, [
        '你是一位擅长制作教学课件的小学【学科】教师。',
        '',
        '请帮我设计一个【年级】【课题】的PPT课件内容大纲：',
        '- 课件页数：约【X】页',
        '- 教学重点：【重点内容】',
        '',
        '请按以下格式输出每一页：',
        '',
        '【第X页】标题：______',
        '- 核心内容（3-5个要点）',
        '- 配图建议',
        '- 讲解提示（教师可以说的话）',
        '',
        '要求：',
        '1. 每页内容不超过5个要点',
        '2. 语言适合小学生理解',
        '3. 包含1-2个互动问题'
    ])
    
    # 分页
    doc.add_page_break()
    
    # 5个核心技巧速查
    doc.add_heading('5个核心技巧速查', 1)
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    
    headers = ['技巧', '说明', '示例']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'E8F4F8')
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ['角色设定', '告诉AI它是谁', '「你是一位资深小学语文教师」'],
        ['任务描述', '明确你要什么', '「设计一节三年级《荷花》教案」'],
        ['提供示例', '给AI看范例', '「参考风格：像抢答游戏那样」'],
        ['迭代优化', '追问修改', '「请把导入改成游戏形式」'],
        ['结构化输出', '指定格式', '「请按以下格式输出...」'],
    ]
    for i, row_data in enumerate(data):
        for j, text in enumerate(row_data):
            table.rows[i+1].cells[j].text = text
    
    doc.add_paragraph()
    
    # 推荐AI工具
    doc.add_heading('推荐AI工具', 1)
    
    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Table Grid'
    
    headers2 = ['工具', '网址', '特点']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'E8F4F8')
        cell.paragraphs[0].runs[0].bold = True
    
    data2 = [
        ['Kimi', 'kimi.moonshot.cn', '免费、长文本'],
        ['文心一言', 'yiyan.baidu.com', '中文优化'],
        ['豆包', 'www.doubao.com', '新手友好'],
        ['通义千问', 'tongyi.aliyun.com', '功能全面'],
    ]
    for i, row_data in enumerate(data2):
        for j, text in enumerate(row_data):
            table2.rows[i+1].cells[j].text = text
    
    # 保存文档
    output_path = 'docs/training/ai-prompt-training/prompt-templates.docx'
    doc.save(output_path)
    print(f'提示词模板卡片Word文档已创建：{output_path}')

if __name__ == '__main__':
    create_templates()
